#!/usr/bin/env python3
"""
深度美化 pandoc 生成的 docx 文件。

Pipeline: pandoc → raw.docx → style_docx.py → styled.docx
build_docx.py 自动调用本脚本，无需手动运行。

处理内容：
  - 中文字体：微软雅黑(标题) / 宋体(正文)
  - 英文字体：Calibri(正文) / Cambria(标题)
  - 代码块：Consolas + 浅灰背景 + 左边框
  - 引用块：浅黄背景 + 左边框
  - 表格：中蓝表头(白字,含数学符号) + 斑马纹
  - 标题层级颜色体系 + H1 底线
  - 页眉(文档标题) + 页脚(页码)
  - 图片最大宽度限制

Usage:
    python3 style_docx.py --input INPUT_raw.docx --output OUTPUT.docx --fallback-title 文档标题
"""

from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

# ── OMML namespace (pandoc converts $...$ to OMML math runs) ──────────────────
NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ═══════════════════════════════════════════════════════════════════════════════
# 颜色体系
# ═══════════════════════════════════════════════════════════════════════════════
COLOR_H1 = RGBColor(0x1A, 0x3C, 0x6E)       # 深蓝
COLOR_H2 = RGBColor(0x2E, 0x5C, 0x8A)       # 中蓝
COLOR_H3 = RGBColor(0x3A, 0x7C, 0xA5)       # 浅蓝
COLOR_H4 = RGBColor(0x55, 0x6B, 0x2E)       # 墨绿
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)     # 近黑
COLOR_QUOTE = RGBColor(0x55, 0x55, 0x55)    # 中灰
COLOR_TABLE_HEADER_BG = "2C5282"            # 中蓝表头（仅用于非表头场景，如 H1 底线）
COLOR_TABLE_HEADER_BG_LIGHT = "D6E4F0"      # 浅蓝表头（OMML 数学符号兼容）
COLOR_TABLE_ZEBRA = "EDF2F7"                # 极浅蓝斑马纹
COLOR_CODE_BG = "F5F5F5"                    # 代码浅灰背景
COLOR_QUOTE_BG = "FFF9E6"                   # 引用浅黄背景
COLOR_CAPTION_TITLE = RGBColor(0x66, 0x66, 0x66)  # 图标题灰
COLOR_FIGURE_NOTE = RGBColor(0x44, 0x44, 0x44)    # 图注深灰

# ═══════════════════════════════════════════════════════════════════════════════
# 字体
# ═══════════════════════════════════════════════════════════════════════════════
FONT_CN_HEADING = "微软雅黑"
FONT_CN_BODY = "宋体"
FONT_EN_HEADING = "Cambria"
FONT_EN_BODY = "Calibri"
FONT_CODE = "Consolas"

DEFAULT_FALLBACK_TITLE = "AI数学：从起步到前沿"

# ═══════════════════════════════════════════════════════════════════════════════
# 低层 XML 工具
# ═══════════════════════════════════════════════════════════════════════════════

def _ensure_child(parent, tag: str):
    """找到或创建 parent 下的指定 XML 子元素。"""
    elem = parent.find(qn(tag))
    if elem is None:
        elem = parse_xml(f'<{tag} {nsdecls("w")}/>')
        parent.append(elem)
    return elem


def _ensure_rpr(run_element):
    """找到或创建 run 的 <w:rPr>。"""
    rpr = run_element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run_element.insert(0, rpr)
    return rpr


def _set_xml_color(rpr, color_hex: str):
    """在 <w:rPr> 中强制设置 <w:color w:val=...>，覆盖 pandoc 默认值。"""
    color_elem = rpr.find(qn('w:color'))
    if color_elem is None:
        color_elem = parse_xml(f'<w:color {nsdecls("w")}/>')
        rpr.append(color_elem)
    color_elem.set(qn('w:val'), color_hex)


def _set_math_run_color(m_run, color_hex: str):
    """给 OMML 数学 run (<m:r>) 设置文字颜色。

    pandoc 把表头里的 $w_1$ 等内联数学转成 OMML <m:r>，
    不受 python-docx 的 run.font.color 控制，必须直接操作 XML。
    """
    mrpr = m_run.find(f'{{{NS_M}}}rPr')
    if mrpr is None:
        mrpr = etree.SubElement(m_run, f'{{{NS_M}}}rPr')
        m_run.insert(0, mrpr)
    # 清除已有颜色
    for old in mrpr.findall(f'{{{NS_W}}}color'):
        mrpr.remove(old)
    color_elem = etree.SubElement(mrpr, f'{{{NS_W}}}color')
    color_elem.set(qn('w:val'), color_hex)


def _set_paragraph_shading(paragraph, color_hex: str):
    """给段落加背景色。"""
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    pPr.append(shading)


def _set_cell_shading(cell, color_hex: str):
    """给单元格加背景色。"""
    tcPr = cell._element.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def _set_paragraph_spacing(paragraph, before=None, after=None, line=None):
    """设置段落间距。"""
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _add_bottom_border(paragraph, color=COLOR_TABLE_HEADER_BG, size="6"):
    """给段落底部加横线。"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_left_border(paragraph, color="3A7CA5", size="18"):
    """给段落左侧加竖线（引用块用）。"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="8" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# Run / 段落样式工具
# ═══════════════════════════════════════════════════════════════════════════════

def set_run_font(run, cn_font, en_font, size=None, color=None, bold=None):
    """设置 run 的字体（中文+英文分别设置）+ 颜色 + 粗体。

    颜色在 XML 层面强制写入 <w:color>，覆盖 pandoc 的默认样式。
    """
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = en_font

    rpr = _ensure_rpr(run._element)
    rfonts = _ensure_child(rpr, 'w:rFonts')
    rfonts.set(qn('w:eastAsia'), cn_font)
    rfonts.set(qn('w:ascii'), en_font)
    rfonts.set(qn('w:hAnsi'), en_font)

    if color is not None:
        run.font.color.rgb = color  # python-docx 高层 API
        _set_xml_color(rpr, str(color))  # XML 层面强制覆盖


def style_runs_in_paragraph(para, cn_font, en_font, size=None, color=None, bold=None):
    """对段落内所有 <w:r> run 应用字体样式。"""
    for run in para.runs:
        set_run_font(run, cn_font, en_font, size=size, color=color, bold=bold)


def style_math_runs_in_paragraph(para, color_hex: str):
    """对段落内所有 OMML <m:r> 数学 run 设置颜色。"""
    for m_run in para._element.iter(f'{{{NS_M}}}r'):
        _set_math_run_color(m_run, color_hex)


def style_all_runs_in_paragraph(para, cn_font, en_font, size=None, color=None, bold=None):
    """对段落内所有 run（普通 + 数学）应用统一样式。"""
    color_hex = str(color) if color is not None else None
    style_runs_in_paragraph(para, cn_font, en_font, size=size, color=color, bold=bold)
    if color_hex:
        style_math_runs_in_paragraph(para, color_hex)


# ═══════════════════════════════════════════════════════════════════════════════
# 文档级样式设置
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_style_font(style, cn_font, en_font):
    """给一个样式对象设置中英文字体。"""
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.insert(0, rpr)
    rfonts = _ensure_child(rpr, 'w:rFonts')
    rfonts.set(qn('w:eastAsia'), cn_font)
    rfonts.set(qn('w:ascii'), en_font)
    rfonts.set(qn('w:hAnsi'), en_font)


def setup_page_layout(doc):
    """页面边距。"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)


def setup_default_styles(doc):
    """Normal 样式 + Heading 1-4 样式 + 超链接样式。"""
    # --- Normal (正文默认) ---
    normal = doc.styles['Normal']
    normal.font.name = FONT_EN_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BODY
    _apply_style_font(normal, FONT_CN_BODY, FONT_EN_BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    # --- Headings ---
    heading_configs = {
        'Heading 1': (20, COLOR_H1, 24, 12),
        'Heading 2': (16, COLOR_H2, 18, 10),
        'Heading 3': (13, COLOR_H3, 14, 8),
        'Heading 4': (12, COLOR_H4, 10, 6),
    }
    for name, (size, color, before, after) in heading_configs.items():
        try:
            style = doc.styles[name]
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
            style.font.name = FONT_EN_HEADING
            _apply_style_font(style, FONT_CN_HEADING, FONT_EN_HEADING)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
        except KeyError:
            pass

    # --- Hyperlink ---
    try:
        hl = doc.styles['Hyperlink']
        hl_rpr = hl.element.find(qn('w:rPr'))
        if hl_rpr is None:
            hl_rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            hl.element.append(hl_rpr)
        for tag in ('w:color', 'w:u'):
            existing = hl_rpr.find(qn(tag))
            if existing is not None:
                hl_rpr.remove(existing)
        hl_rpr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="0563C1"/>'))
        hl_rpr.append(parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>'))
    except KeyError:
        pass


# ═══════════════════════════════════════════════════════════════════════════════
# 段落处理
# ═══════════════════════════════════════════════════════════════════════════════

def process_paragraphs(doc):
    """逐段落应用样式：代码块、标题、引用、图注、正文。"""

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        # 代码块
        if style_name == 'Source Code':
            _set_paragraph_shading(para, COLOR_CODE_BG)
            _set_paragraph_spacing(para, before=2, after=2, line=1.15)
            style_runs_in_paragraph(para, FONT_CODE, FONT_CODE,
                                    size=9.5, color=RGBColor(0x33, 0x33, 0x33))
            _add_left_border(para, color="CCCCCC", size="12")
            continue

        # 标题
        if style_name == 'Heading 1':
            style_runs_in_paragraph(para, FONT_CN_HEADING, FONT_EN_HEADING,
                                    size=20, color=COLOR_H1, bold=True)
            _add_bottom_border(para, color=COLOR_TABLE_HEADER_BG, size="8")
            continue

        if style_name == 'Heading 2':
            style_runs_in_paragraph(para, FONT_CN_HEADING, FONT_EN_HEADING,
                                    size=16, color=COLOR_H2, bold=True)
            continue

        if style_name == 'Heading 3':
            style_runs_in_paragraph(para, FONT_CN_HEADING, FONT_EN_HEADING,
                                    size=13, color=COLOR_H3, bold=True)
            if '自检答案' in text:
                pPr = para._element.get_or_add_pPr()
                pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))
            continue

        if style_name == 'Heading 4':
            style_runs_in_paragraph(para, FONT_CN_HEADING, FONT_EN_HEADING,
                                    size=12, color=COLOR_H4, bold=True)
            continue

        # 引用块
        if style_name in ('Block Text', 'Quote', 'Intense Quote'):
            _set_paragraph_shading(para, COLOR_QUOTE_BG)
            _set_paragraph_spacing(para, before=6, after=6, line=1.3)
            _add_left_border(para, color="E0A800", size="18")
            style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                    size=10.5, color=COLOR_QUOTE)
            continue

        # 自检答案段落 (A1：... A2：...)
        if text and re.match(r'^A\d+：', text):
            _set_paragraph_shading(para, "F0FFF4")
            _set_paragraph_spacing(para, before=4, after=8, line=1.3)
            _add_left_border(para, color="38A169", size="14")
            style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                    size=10.5, color=COLOR_BODY)
            continue

        # 图片段落
        if style_name == 'Captioned Figure':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 图片标题
        if style_name == 'Image Caption':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(para, before=4, after=2, line=1.15)
            style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                    size=9.5, color=COLOR_CAPTION_TITLE, bold=True)
            continue

        # 图注段落 ("图1：..." 或 "图1:...")
        if style_name == 'Body Text' and re.match(r'^图\d+[：:]', text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(para, before=2, after=12, line=1.35)
            style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                    size=9.5, color=COLOR_FIGURE_NOTE)
            continue

        # 正文段落
        if style_name in ('Normal', 'Body Text', 'First Paragraph', 'Compact') and text:
            style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                    size=11, color=COLOR_BODY)


# ═══════════════════════════════════════════════════════════════════════════════
# 表格处理
# ═══════════════════════════════════════════════════════════════════════════════

def _style_table_header_cell(cell):
    """表头单元格：浅蓝底 + 深字(含数学符号) + 居中。

    使用浅色背景避免 Word OMML 数学引擎的自动反色：
    当背景为深色时，Word 会将白色数学符号自动转为黑色，
    导致在蓝色表头上看不清。浅色背景 + 深色字体可可靠地
    避免此问题。
    """
    _set_cell_shading(cell, COLOR_TABLE_HEADER_BG_LIGHT)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 普通文本 run 无文字但有 text 时重建 run（pandoc 偶尔生成空 run 段落）
        if not para.runs and para.text.strip():
            for child in list(para._element):
                if child.tag == qn('w:r'):
                    para._element.remove(child)
            para.add_run(para.text)
        # 深字 + 粗体（普通 run + 数学 run）
        style_all_runs_in_paragraph(para, FONT_CN_HEADING, FONT_EN_HEADING,
                                    size=10.5, color=COLOR_BODY, bold=True)


def _style_table_body_cell(cell, row_idx: int):
    """正文单元格：偶数行斑马纹。"""
    if row_idx % 2 == 0:
        _set_cell_shading(cell, COLOR_TABLE_ZEBRA)
    for para in cell.paragraphs:
        style_runs_in_paragraph(para, FONT_CN_BODY, FONT_EN_BODY,
                                size=10.5, color=COLOR_BODY)


def process_tables(doc):
    """表格美化：边框 + 表头底色白字 + 斑马纹。"""
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表格边框
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblBorders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
                f'  <w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
                f'  <w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:color="E0E0E0"/>'
                f'  <w:insideV w:val="single" w:sz="4" w:color="E0E0E0"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(tblBorders)

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if row_idx == 0:
                    _style_table_header_cell(cell)
                else:
                    _style_table_body_cell(cell, row_idx)


# ═══════════════════════════════════════════════════════════════════════════════
# 页眉页脚
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_numbers(doc):
    """页脚居中页码。"""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.text = ''

        # PAGE field: <fldChar begin/> <instrText> PAGE </instrText> <fldChar end/>
        for fld_type, content in [('begin', None), (None, ' PAGE '), ('end', None)]:
            r = para.add_run()
            if fld_type:
                r._element.append(parse_xml(
                    f'<w:fldChar {nsdecls("w")} w:fldCharType="{fld_type}"/>'
                ))
            else:
                r._element.append(parse_xml(
                    f'<w:instrText {nsdecls("w")} xml:space="preserve">{content}</w:instrText>'
                ))

        for r in para.runs:
            set_run_font(r, FONT_CN_BODY, FONT_EN_BODY, size=9,
                         color=RGBColor(0x99, 0x99, 0x99))


def add_headers(doc, title: str):
    """页眉右对齐文档标题。"""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.text = ''
        run = para.add_run(title)
        set_run_font(run, FONT_CN_HEADING, FONT_EN_HEADING,
                     size=9, color=RGBColor(0xAA, 0xAA, 0xAA))


# ═══════════════════════════════════════════════════════════════════════════════
# 图片宽度限制
# ═══════════════════════════════════════════════════════════════════════════════

def limit_image_widths(doc, max_width_cm: float = 15.0):
    """限制图片最大宽度，等比缩放。"""
    max_width = int(Cm(max_width_cm).emu)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for drawing in run._element.findall(qn('w:drawing')):
                for ext_tag in (qn('wp:extent'), qn('a:ext')):
                    for ext in drawing.findall('.//' + ext_tag):
                        cx = int(ext.get('cx', '0'))
                        cy = int(ext.get('cy', '0'))
                        if cx > max_width:
                            ratio = max_width / cx
                            ext.set('cx', str(max_width))
                            ext.set('cy', str(int(cy * ratio)))


# ═══════════════════════════════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════════════════════════════

def get_document_title(document, fallback=DEFAULT_FALLBACK_TITLE) -> str:
    """优先使用正文第一个 H1 作为页眉标题。"""
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name == 'Heading 1':
            title = paragraph.text.strip()
            if title:
                return title
    return fallback


def style_docx(input_path: str | Path, output_path: str | Path,
               fallback_title: str = DEFAULT_FALLBACK_TITLE,
               cleanup: bool = True) -> None:
    """主函数：读取 pandoc 生成的 raw docx，输出美化后的 styled docx。"""
    input_path = Path(input_path)
    output_path = Path(output_path)
    if not input_path.exists():
        raise FileNotFoundError(f"input docx not found: {input_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))
    title = get_document_title(doc, fallback_title)

    setup_page_layout(doc)
    setup_default_styles(doc)
    process_paragraphs(doc)
    process_tables(doc)
    add_page_numbers(doc)
    add_headers(doc, title)
    limit_image_widths(doc)

    doc.save(str(output_path))

    size = os.path.getsize(output_path)
    print(f"✅ 美化完成: {output_path}")
    print(f"   文件大小: {size / 1024 / 1024:.1f} MB")
    print(f"   段落数: {len(doc.paragraphs)}")
    print(f"   表格数: {len(doc.tables)}")

    if cleanup and input_path.exists():
        input_path.unlink()
        print(f"   已清理中间文件: {input_path}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Apply project DOCX styling to a pandoc-generated .docx file."
    )
    parser.add_argument("--input", "-i", required=True, help="Pandoc raw DOCX input path.")
    parser.add_argument("--output", "-o", required=True, help="Styled DOCX output path.")
    parser.add_argument("--fallback-title", default=DEFAULT_FALLBACK_TITLE,
                        help="Header title if no Heading 1 exists.")
    parser.add_argument("--keep-raw", action="store_true",
                        help="Keep the raw input DOCX instead of deleting it.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    style_docx(args.input, args.output,
               fallback_title=args.fallback_title, cleanup=not args.keep_raw)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
